from openai import OpenAI
from dotenv import load_dotenv
import os
import time
import warnings
import glob
import json
import re
import csv
import PyPDF2
import docx
import sqlite3
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
import pandas as pd
from fpdf import FPDF
import uuid
import openpyxl
import math
import pdfplumber

# =============================================================================
# CONFIGURATION AND INITIALIZATION
# =============================================================================

warnings.filterwarnings("ignore", category=DeprecationWarning)

class AdvancedPDFAnalysisAssistant:
    """
    Advanced document analysis system with multi-format support, batch processing,
    export capabilities, and user authentication.
    """
    
    def __init__(self):
        """Initialize the advanced document analysis system."""
        load_dotenv(dotenv_path="/root/filesearchfolder/.env")
        self.client = OpenAI()
        self.assistant_id = None
        self.vector_store_id = None
        self.file_ids = []
        self.current_user = None
        self.thread_id = None
        self.calculation_mode = "auto"
        self.manual_data = {}  # Store manually entered data
        self.initialize_environment()
        self.init_database()
        self.init_calculation_formulas()

    def initialize_environment(self):
        """Initialize required directories and validate environment."""
        self.base_dir = Path.cwd()
        self.docs_folder = self.base_dir / "documents"
        self.output_folder = self.base_dir / "output"
        self.db_folder = self.base_dir / "database"
        self.temp_folder = self.base_dir / "temp"
        
        # Create necessary directories
        for folder in [self.docs_folder, self.output_folder, self.db_folder, self.temp_folder]:
            folder.mkdir(exist_ok=True)
        
        # Validate API key
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY not found in environment variables")

    def init_database(self):
        """Initialize SQLite database for user authentication and session tracking."""
        self.db_path = self.db_folder / "document_analysis.db"
        self.conn = sqlite3.connect(self.db_path)
        self.cursor = self.conn.cursor()
        
        # Create users table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                email TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                last_login TIMESTAMP
            )
        ''')
        
        # Create sessions table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                session_id TEXT UNIQUE,
                start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                end_time TIMESTAMP,
                queries_executed INTEGER DEFAULT 0,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')
        
        # Create query_history table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS query_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT,
                query_text TEXT,
                response_text TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (session_id) REFERENCES sessions (session_id)
            )
        ''')
        
        self.conn.commit()
        print("✅ Database initialized successfully!")

    def init_calculation_formulas(self):
        """Initialize calculation formulas."""
        self.fixed_formulas = {
            "pension_government": {
                "formula": "50% of average emoluments + 1% for each year beyond 20 years",
                "calculation": lambda basic, da, years: self._fixed_pension_gov(basic, da, years)
            },
            "pension_corporate": {
                "formula": "2% of salary per year of service",
                "calculation": lambda basic, da, years: self._fixed_pension_corporate(basic, da, years)
            },
            "pension_eps": {
                "formula": "(Pensionable Salary * Years of Service) / 70",
                "calculation": lambda basic, da, years: self._fixed_pension_eps(basic, da, years)
            },
            "gratuity_regular": {
                "formula": "(Salary * 15 * Years of Service) / 26",
                "calculation": lambda basic, da, years: self._fixed_gratuity_regular(basic, da, years)
            },
            "gratuity_government": {
                "formula": "Half month salary for each completed 6-month period",
                "calculation": lambda basic, da, years: self._fixed_gratuity_gov(basic, da, years)
            }
        }
        self.extracted_formulas = {}
        print("✅ Calculation formulas initialized!")

    def _fixed_pension_gov(self, basic_salary: float, da: float, years_of_service: int) -> float:
        """Fixed government pension calculation."""
        total_salary = basic_salary + da
        if years_of_service >= 10:
            pension_amount = 0.5 * total_salary
            if years_of_service > 20:
                pension_amount += (years_of_service - 20) * 0.01 * total_salary
            return min(pension_amount, total_salary)
        return 0

    def _fixed_pension_corporate(self, basic_salary: float, da: float, years_of_service: int) -> float:
        """Fixed corporate pension calculation."""
        total_salary = basic_salary + da
        pension_amount = (total_salary * years_of_service * 0.02)
        return min(pension_amount, total_salary * 0.5)

    def _fixed_pension_eps(self, basic_salary: float, da: float, years_of_service: int) -> float:
        """Fixed EPS pension calculation."""
        total_salary = basic_salary + da
        pensionable_salary = min(total_salary, 15000)
        return (pensionable_salary * years_of_service) / 70

    def _fixed_gratuity_regular(self, basic_salary: float, da: float, years_of_service: float) -> float:
        """Fixed regular gratuity calculation."""
        total_salary = basic_salary + da
        if years_of_service >= 5:
            gratuity_amount = (total_salary * 15 * years_of_service) / 26
            return min(gratuity_amount, 2000000)
        return 0

    def _fixed_gratuity_gov(self, basic_salary: float, da: float, years_of_service: float) -> float:
        """Fixed government gratuity calculation."""
        total_salary = basic_salary + da
        return (total_salary * years_of_service) / 2

    def hash_password(self, password: str) -> str:
        """Hash password using SHA-256."""
        return hashlib.sha256(password.encode()).hexdigest()

    def register_user(self, username: str, password: str, email: str = None) -> bool:
        """Register a new user."""
        try:
            password_hash = self.hash_password(password)
            self.cursor.execute(
                "INSERT INTO users (username, password_hash, email) VALUES (?, ?, ?)",
                (username, password_hash, email)
            )
            self.conn.commit()
            print(f"✅ User '{username}' registered successfully!")
            return True
        except sqlite3.IntegrityError:
            print("❌ Username already exists")
            return False
        except Exception as e:
            print(f"❌ Registration failed: {e}")
            return False

    def login_user(self, username: str, password: str) -> bool:
        """Authenticate user and start session."""
        try:
            password_hash = self.hash_password(password)
            self.cursor.execute(
                "SELECT id FROM users WHERE username = ? AND password_hash = ?",
                (username, password_hash)
            )
            user = self.cursor.fetchone()
            
            if user:
                self.current_user = {
                    'id': user[0],
                    'username': username
                }
                
                session_id = str(uuid.uuid4())
                self.cursor.execute(
                    "INSERT INTO sessions (user_id, session_id) VALUES (?, ?)",
                    (user[0], session_id)
                )
                self.conn.commit()
                
                self.current_user['session_id'] = session_id
                print(f"✅ Welcome back, {username}!")
                return True
            else:
                print("❌ Invalid username or password")
                return False
                
        except Exception as e:
            print(f"❌ Login failed: {e}")
            return False

    def logout_user(self):
        """End current user session."""
        if self.current_user and 'session_id' in self.current_user:
            self.cursor.execute(
                "UPDATE sessions SET end_time = CURRENT_TIMESTAMP WHERE session_id = ?",
                (self.current_user['session_id'],)
            )
            self.conn.commit()
            print(f"👋 Goodbye, {self.current_user['username']}!")
        self.current_user = None

    def authenticate_user(self) -> bool:
        """Handle user authentication with multiple options."""
        print("🔐 Document Analysis System - Authentication")
        print("="*50)
        while True:
            print("\nOptions:")
            print("1. Login")
            print("2. Register")
            print("3. Continue as Guest")
            print("4. Exit")
            choice = input("Select option (1-4): ").strip()
            if choice == '1':
                username = input("Username: ").strip()
                password = input("Password: ").strip()
                if self.login_user(username, password):
                    return True
            elif choice == '2':
                username = input("New username: ").strip()
                password = input("New password: ").strip()
                email = input("Email (optional): ").strip() or None
                if self.register_user(username, password, email):
                    if self.login_user(username, password):
                        return True
            elif choice == '3':
                print("ℹ️  Continuing as guest user (limited functionality)")
                return True
            elif choice == '4':
                print("👋 Goodbye!")
                return False
            else:
                print("❌ Invalid option")

    def get_supported_files(self, folder_path: Path) -> List[Path]:
        """Get all supported document files."""
        all_files = list(folder_path.iterdir())
        supported_extensions = ['.pdf', '.docx', '.txt', '.csv', '.xlsx']
        files = [
            file for file in all_files 
            if file.is_file() and file.suffix.lower() in supported_extensions
        ]
        return self._validate_files(files)

    def _validate_files(self, files: List[Path]) -> List[Path]:
        """Validate files for size and format requirements."""
        valid_files = []
        max_size_mb = 20
        
        for file_path in files:
            try:
                file_size_mb = file_path.stat().st_size / (1024 * 1024)
                if file_size_mb > max_size_mb:
                    print(f"⚠️  Skipping {file_path.name} - Too large ({file_size_mb:.1f}MB)")
                    continue
                valid_files.append(file_path)
            except Exception as e:
                print(f"⚠️  Skipping {file_path.name} - Error: {e}")
        return valid_files

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file formats with improved PDF handling."""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_text_from_pdf_enhanced(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_path.suffix.lower() in ['.txt', '.csv']:
                return self._extract_text_from_text_file(file_path)
            elif file_path.suffix.lower() == '.xlsx':
                return self._extract_text_from_excel(file_path)
            else:
                return ""
        except Exception as e:
            print(f"❌ Error extracting text from {file_path.name}: {e}")
            return ""

    def _extract_text_from_pdf_enhanced(self, file_path: Path) -> str:
        """Enhanced PDF text extraction using multiple methods."""
        text = ""
        
        print(f"🔍 Extracting text from {file_path.name}...")
        
        # Method 1: Try pdfplumber first
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text and len(page_text.strip()) > 10:
                        text += f"--- Page {i+1} ---\n{page_text}\n"
                        print(f"✅ Page {i+1}: Extracted {len(page_text)} characters")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = f"--- Table {table_idx+1} on Page {i+1} ---\n"
                            for row in table:
                                if row and any(cell is not None for cell in row):
                                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                    table_text += row_text + "\n"
                            if len(table_text) > 50:
                                text += table_text + "\n"
            
            if text and len(text.strip()) > 50:
                print(f"✅ Successfully extracted {len(text)} characters using pdfplumber")
                return text
            else:
                print("⚠️ pdfplumber extracted very little text, trying PyPDF2...")
        except Exception as e:
            print(f"⚠️ pdfplumber failed: {e}")
        
        # Method 2: Try PyPDF2
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and len(page_text.strip()) > 10:
                        text += f"--- Page {i+1} ---\n{page_text}\n"
                        print(f"✅ Page {i+1}: Extracted {len(page_text)} characters")
            
            if text and len(text.strip()) > 50:
                print(f"✅ Successfully extracted {len(text)} characters using PyPDF2")
                return text
            else:
                print("⚠️ Both PDF libraries extracted very little text - document is likely image-based")
        except Exception as e:
            print(f"⚠️ PyPDF2 failed: {e}")
        
        return text

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Also extract tables from DOCX
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            return ""

    def _extract_text_from_text_file(self, file_path: Path) -> str:
        """Extract text from text-based files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"❌ Text file extraction error: {e}")
            return ""

    def _extract_text_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel files."""
        try:
            df = pd.read_excel(file_path)
            text_representation = ""
            
            # Add column names
            text_representation += "Columns: " + ", ".join(df.columns.astype(str)) + "\n\n"
            
            # Add sample data (first 10 rows)
            for idx, row in df.head(10).iterrows():
                row_text = " | ".join([str(val) for val in row.values])
                text_representation += f"Row {idx}: {row_text}\n"
            
            return text_representation
        except Exception as e:
            print(f"❌ Excel extraction error: {e}")
            return ""

    def offer_manual_data_entry(self, filename: str) -> bool:
        """Offer manual data entry for image-based PDFs."""
        print(f"\n📝 MANUAL DATA ENTRY for {filename}")
        print("Since this is an image-based PDF, you can manually enter the financial data.")
        print("This data will be stored and used for calculations.")
        
        choice = input("Would you like to enter salary data manually? (y/n): ").strip().lower()
        if choice in ['y', 'yes']:
            return self._collect_manual_salary_data(filename)
        return False

    def _collect_manual_salary_data(self, filename: str) -> bool:
        """Collect salary data manually from user."""
        try:
            print("\n💰 Please enter the salary details:")
            
            data = {}
            
            # Basic salary information
            basic = input("Basic Salary: ").strip()
            if basic:
                data['basic_salary'] = float(basic.replace(',', ''))
            
            da = input("DA (Dearness Allowance): ").strip()
            if da:
                data['da'] = float(da.replace(',', ''))
            
            hra = input("HRA (House Rent Allowance): ").strip()
            if hra:
                data['hra'] = float(hra.replace(',', ''))
            
            # Gross and Net
            gross = input("Gross Salary: ").strip()
            if gross:
                data['gross_salary'] = float(gross.replace(',', ''))
            
            net = input("Net Salary (Take Home): ").strip()
            if net:
                data['net_salary'] = float(net.replace(',', ''))
            
            # Deductions
            pf = input("PF Deduction: ").strip()
            if pf:
                data['pf_deduction'] = float(pf.replace(',', ''))
            
            tax = input("Tax Deduction: ").strip()
            if tax:
                data['tax_deduction'] = float(tax.replace(',', ''))
            
            other_deductions = input("Other Deductions: ").strip()
            if other_deductions:
                data['other_deductions'] = float(other_deductions.replace(',', ''))
            
            total_deductions = input("Total Deductions: ").strip()
            if total_deductions:
                data['total_deductions'] = float(total_deductions.replace(',', ''))
            
            # Store the data
            if data:
                self.manual_data[filename] = data
                print(f"✅ Successfully stored manual data for {filename}")
                print(f"📊 Data entered: {data}")
                return True
            else:
                print("❌ No data entered")
                return False
                
        except ValueError as e:
            print(f"❌ Invalid number format: {e}")
            return False
        except Exception as e:
            print(f"❌ Error collecting manual data: {e}")
            return False

    def get_manual_data_summary(self, filename: str) -> str:
        """Get summary of manually entered data."""
        if filename in self.manual_data:
            data = self.manual_data[filename]
            summary = "## 📝 Manually Entered Salary Data\n\n"
            summary += f"**Document:** {filename}\n\n"
            summary += "**Salary Breakdown:**\n"
            
            for key, value in data.items():
                display_key = key.replace('_', ' ').title()
                summary += f"- **{display_key}:** ₹{value:,.2f}\n"
            
            # Calculate any missing values
            if 'gross_salary' in data and 'net_salary' in data and 'total_deductions' not in data:
                total_deductions = data['gross_salary'] - data['net_salary']
                summary += f"- **Calculated Total Deductions:** ₹{total_deductions:,.2f}\n"
            
            summary += f"\n*Data entered manually on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*"
            return summary
        else:
            return ""

    def process_documents_in_batches(self, batch_size: int = 10) -> List[str]:
        """Process documents in batches with enhanced error handling and manual entry options."""
        all_files = self.get_supported_files(self.docs_folder)
        if not all_files:
            print("ℹ️  No documents found for processing")
            return []
        
        print(f"📦 Found {len(all_files)} documents. Processing in batches of {batch_size}...")
        processed_files = []
        
        # First, analyze each file
        print("\n🔍 Detailed Document Analysis:")
        for file_path in all_files:
            print(f"\n📄 Analyzing: {file_path.name}")
            text_content = self.extract_text_from_file(file_path)
            
            if text_content and len(text_content.strip()) > 10:
                content_length = len(text_content)
                print(f"   ✅ Text extracted: {content_length} characters")
                
                # Show sample if meaningful
                if content_length > 50:
                    sample = text_content[:300].replace('\n', ' ')
                    sample = re.sub(r'\s+', ' ', sample)
                    print(f"   📄 Sample: {sample}...")
            else:
                print(f"   ❌ IMAGE-BASED PDF - No text extracted")
                # Offer manual data entry for image-based PDFs
                if file_path.suffix.lower() == '.pdf':
                    self.offer_manual_data_entry(file_path.name)
        
        # Process files for upload
        for i in range(0, len(all_files), batch_size):
            batch = all_files[i:i + batch_size]
            print(f"\n🔄 Processing batch {i//batch_size + 1}/{(len(all_files)-1)//batch_size + 1}")
            try:
                batch_file_ids = self._process_batch(batch)
                processed_files.extend(batch_file_ids)
                print(f"✅ Batch processed successfully. Total files: {len(processed_files)}")
                if i + batch_size < len(all_files):
                    time.sleep(5)
            except Exception as e:
                print(f"❌ Batch processing failed: {e}")
                continue
        return processed_files

    def _process_batch(self, batch_files: List[Path]) -> List[str]:
        """Process a single batch of files."""
        batch_file_ids = []
        processed_filenames = set()
        
        for file_path in batch_files:
            try:
                if file_path.name.lower() in processed_filenames:
                    print(f"⚠️  Skipping duplicate: {file_path.name}")
                    continue
                    
                # Extract text
                text_content = self.extract_text_from_file(file_path)
                
                # Create content for upload
                if text_content and len(text_content.strip()) > 10:
                    # Create enhanced text file
                    enhanced_content = f"""
DOCUMENT: {file_path.name}
STATUS: Text extracted successfully
CONTENT LENGTH: {len(text_content)} characters

EXTRACTED CONTENT:
{text_content}
"""
                    temp_file = self.temp_folder / f"{file_path.stem}_enhanced.txt"
                    with open(temp_file, 'w', encoding='utf-8') as f:
                        f.write(enhanced_content)
                    
                    with open(temp_file, 'rb') as f:
                        file_obj = self.client.files.create(file=f, purpose="assistants")
                        batch_file_ids.append(file_obj.id)
                        processed_filenames.add(file_path.name.lower())
                        print(f"✅ Uploaded: {file_path.name} ({len(text_content)} chars)")
                    
                    temp_file.unlink()
                else:
                    # For image-based PDFs, create a helpful error file
                    error_content = f"""
DOCUMENT: {file_path.name}
STATUS: IMAGE-BASED PDF - No text could be extracted

SOLUTIONS:
1. Use OCR software to convert this PDF to text
2. Use the manual data entry feature by typing 'manual'
3. Provide a text-based PDF instead of scanned/image PDF

MANUAL DATA AVAILABLE: {'Yes' if file_path.name in self.manual_data else 'No'}
"""
                    temp_file = self.temp_folder / f"{file_path.stem}_error.txt"
                    with open(temp_file, 'w', encoding='utf-8') as f:
                        f.write(error_content)
                    
                    with open(temp_file, 'rb') as f:
                        file_obj = self.client.files.create(file=f, purpose="assistants")
                        batch_file_ids.append(file_obj.id)
                        processed_filenames.add(file_path.name.lower())
                        print(f"⚠️  Uploaded error report: {file_path.name}")
                    
                    temp_file.unlink()
                    
            except Exception as e:
                print(f"❌ Failed to process {file_path.name}: {e}")
        return batch_file_ids

    def wait_for_vector_store_processing(self, vector_store_id: str, max_attempts: int = 30, delay: int = 5):
        """Wait for vector store processing to complete."""
        print("⏳ Waiting for vector store processing...")
        for attempt in range(max_attempts):
            try:
                vector_store = self.client.vector_stores.retrieve(vector_store_id)
                status = vector_store.status
                if status == "completed":
                    print("✅ Vector store processing completed!")
                    return True
                elif status == "failed":
                    print("❌ Vector store processing failed!")
                    return False
                elif status in ["in_progress", "cancelling"]:
                    print(f"🔄 Processing... (Attempt {attempt + 1}/{max_attempts})")
                    time.sleep(delay)
                else:
                    print(f"ℹ️  Unknown vector store status: {status}")
                    time.sleep(delay)
            except Exception as e:
                print(f"⚠️  Error checking vector store status: {e}")
                time.sleep(delay)
        print("❌ Vector store processing timed out!")
        return False

    def setup_assistant(self, file_ids: List[str]):
        """Set up the assistant with the provided file IDs."""
        try:
            print("🏗️  Creating vector store...")
            vector_store = self.client.vector_stores.create(
                name=f"Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                file_ids=file_ids
            )
            self.vector_store_id = vector_store.id
            
            if not self.wait_for_vector_store_processing(self.vector_store_id):
                print("⚠️  Vector store processing had issues, but continuing...")
            
            print("👨‍💼 Creating assistant...")
            assistant = self.client.beta.assistants.create(
                name="Advanced Document Analysis Assistant",
                instructions="""You are an expert document analysis assistant. Analyze ALL uploaded documents thoroughly.

SPECIAL INSTRUCTIONS FOR FINANCIAL DOCUMENTS:
1. If documents contain payslips or financial statements, extract ALL numbers and amounts
2. Look for: Gross Salary, Net Salary, Basic Pay, HRA, DA, PF, Taxes, Deductions
3. Provide exact amounts and complete breakdowns

IMPORTANT: Some documents may be image-based with no extractable text. In such cases, 
rely on any manually provided data or indicate that the document needs manual data entry.""",
                model="gpt-4o",
                tools=[{"type": "file_search"}],
                tool_resources={"file_search": {"vector_store_ids": [self.vector_store_id]}}
            )
            self.assistant_id = assistant.id
            
            # Create thread here to avoid None thread_id errors
            thread = self.client.beta.threads.create()
            self.thread_id = thread.id
            print("✅ Assistant setup completed successfully!")
        except Exception as e:
            print(f"❌ Assistant setup failed: {e}")
            raise

    def execute_query(self, query: str):
        """Execute a query against the assistant."""
        print(f"🤔 Asking: {query}")
        try:
            # Ensure we have a thread
            if not self.thread_id:
                thread = self.client.beta.threads.create()
                self.thread_id = thread.id
            
            self.client.beta.threads.messages.create(
                thread_id=self.thread_id,
                role="user",
                content=query
            )
            
            run = self.client.beta.threads.runs.create(
                thread_id=self.thread_id,
                assistant_id=self.assistant_id
            )
            
            print("⏳ Processing your query...")
            while run.status in ['queued', 'in_progress']:
                time.sleep(2)
                run = self.client.beta.threads.runs.retrieve(
                    thread_id=self.thread_id,
                    run_id=run.id
                )
            
            if run.status == 'completed':
                messages = self.client.beta.threads.messages.list(thread_id=self.thread_id)
                assistant_messages = [msg for msg in messages.data if msg.role == 'assistant']
                
                if assistant_messages:
                    response = assistant_messages[0].content[0].text.value
                    print(f"🤖 Assistant: {response}")
                    
                    if self.current_user and 'session_id' in self.current_user:
                        self.cursor.execute(
                            "INSERT INTO query_history (session_id, query_text, response_text) VALUES (?, ?, ?)",
                            (self.current_user['session_id'], query, response)
                        )
                        self.conn.commit()
                    return response
                else:
                    print("❌ No response from assistant")
                    return None
            else:
                print(f"❌ Run failed with status: {run.status}")
                if run.last_error:
                    print(f"Error: {run.last_error.message}")
                return None
        except Exception as e:
            print(f"❌ Error executing query: {e}")
            return None

    def _process_financial_query(self, query: str) -> str:
        """Process financial-related queries with manual data support."""
        # First check for manual data
        all_files = self.get_supported_files(self.docs_folder)
        manual_results = ""
        
        for file_path in all_files:
            manual_summary = self.get_manual_data_summary(file_path.name)
            if manual_summary:
                manual_results += manual_summary + "\n\n"
        
        if manual_results:
            print("📝 Using manually entered data...")
        
        # Then do AI search
        print("🔍 Performing AI document search...")
        
        financial_search_query = f"""
USER QUERY: "{query}"

SEARCH INSTRUCTIONS:
1. First, check if there are any manually entered financial data in the documents
2. Then search all document content for financial information
3. Look for: Gross Salary, Net Salary, Basic Pay, HRA, DA, PF, Taxes, Deductions
4. Provide exact amounts and complete breakdowns

IMPORTANT: Some documents may be image-based with no extractable text. In such cases, 
rely on any manually provided data or indicate that the document needs manual data entry.
"""
        
        ai_response = self.execute_query(financial_search_query)
        
        if manual_results:
            return f"{manual_results}## AI Document Analysis:\n{ai_response}"
        else:
            return ai_response

    def _process_special_query(self, query: str, filename: str) -> str:
        """Process queries with enhanced document analysis."""
        try:
            query_lower = query.lower()
            
            # Check if this is a financial query
            financial_terms = ['salary', 'gross', 'net', 'deduction', 'payslip', 'earnings', 
                             'basic', 'hra', 'da', 'pf', 'tax', 'income']
            if any(term in query_lower for term in financial_terms):
                print(f"💰 Processing financial query: {query}")
                return self._process_financial_query(query)
            
            # Enhanced document search with better context
            print(f"🔍 Searching across ALL documents for: {query}")
            
            # Get list of all documents for context
            all_files = self.get_supported_files(self.docs_folder)
            file_list = [f.name for f in all_files]
            
            enhanced_query = f"""
Please search across ALL available documents to answer this query: "{query}"

Available documents: {', '.join(file_list)}

Search through all these documents comprehensively to find the most relevant information.
If the information exists in any of the documents, provide a complete answer.
If you find conflicting information across documents, mention the differences.
If no information is found in any document, please state that clearly.

Provide specific references to which document contained the information when possible.
"""
            response = self.execute_query(enhanced_query)
            return response if response else "No response received from the assistant"
            
        except Exception as e:
            return f"Error processing query: {str(e)}"

    def set_calculation_mode(self, mode: str):
        """Set calculation mode: auto, document, or fixed."""
        valid_modes = ["auto", "document", "fixed"]
        if mode.lower() in valid_modes:
            self.calculation_mode = mode.lower()
            print(f"✅ Calculation mode set to: {self.calculation_mode}")
        else:
            print(f"❌ Invalid mode. Use: {', '.join(valid_modes)}")

    def interactive_mode(self):
        """Enter interactive mode with enhanced options."""
        print("\n" + "="*60)
        print("💬 INTERACTIVE MODE - Ask questions about your documents")
        print("Type 'exit' to quit, 'manual' to enter data, or 'help' for options")
        print("="*60)
        
        # Show document status
        all_files = self.get_supported_files(self.docs_folder)
        print(f"\n📋 Loaded Documents: {len(all_files)} file(s)")
        for file_path in all_files:
            # Quick text extraction check
            text_content = self.extract_text_from_file(file_path)
            if text_content and len(text_content) > 100:
                status = "✅ Text extracted"
            elif file_path.name in self.manual_data:
                status = "📝 Manual data available"
            else:
                status = "❌ Image-based (needs manual entry)"
            print(f"   - {file_path.name}: {status}")
        
        query_count = 0
        query_history = []
        
        while True:
            user_input = input("\n🎯 Your question: ").strip()
            if user_input.lower() in ['exit', 'quit', 'bye']:
                print("Exiting interactive mode...")
                break
            elif user_input.lower() in ['help', '?']:
                self.show_help()
                continue
            elif user_input.lower() == 'export':
                self.handle_export(query_history)
                continue
            elif user_input.lower() == 'manual':
                self.handle_manual_data_entry()
                continue
            elif user_input.lower() == 'status':
                self.show_document_status()
                continue
            elif not user_input:
                continue
            
            response = self._process_special_query(user_input, "")
            if response:
                query_count += 1
                query_history.append({
                    'query': user_input,
                    'response': response,
                    'timestamp': datetime.now().isoformat()
                })

    def handle_manual_data_entry(self):
        """Handle manual data entry for all documents."""
        all_files = self.get_supported_files(self.docs_folder)
        if not all_files:
            print("❌ No documents available for data entry")
            return
        
        print("\n📝 MANUAL DATA ENTRY")
        print("====================")
        print("Select a document to enter data for:")
        
        for i, file_path in enumerate(all_files, 1):
            has_data = "📊" if file_path.name in self.manual_data else "📝"
            print(f"{i}. {has_data} {file_path.name}")
        
        print(f"{len(all_files) + 1}. View all entered data")
        
        try:
            choice = input(f"\nSelect option (1-{len(all_files) + 1}): ").strip()
            if choice.isdigit():
                choice_num = int(choice)
                if 1 <= choice_num <= len(all_files):
                    selected_file = all_files[choice_num - 1]
                    self._collect_manual_salary_data(selected_file.name)
                elif choice_num == len(all_files) + 1:
                    self.show_all_manual_data()
                else:
                    print("❌ Invalid selection")
            else:
                print("❌ Please enter a number")
        except Exception as e:
            print(f"❌ Error: {e}")

    def show_all_manual_data(self):
        """Show all manually entered data."""
        if not self.manual_data:
            print("❌ No manual data entered yet")
            return
        
        print("\n📊 ALL MANUALLY ENTERED DATA")
        print("============================")
        for filename, data in self.manual_data.items():
            print(f"\n📄 {filename}:")
            for key, value in data.items():
                display_key = key.replace('_', ' ').title()
                print(f"   - {display_key}: ₹{value:,.2f}")

    def show_document_status(self):
        """Show detailed status of loaded documents."""
        print("\n📊 DOCUMENT STATUS REPORT")
        print("="*50)
        
        all_files = self.get_supported_files(self.docs_folder)
        if not all_files:
            print("No documents loaded")
            return
        
        for file_path in all_files:
            print(f"\n📄 {file_path.name}:")
            text_content = self.extract_text_from_file(file_path)
            
            if text_content and len(text_content.strip()) > 10:
                content_length = len(text_content)
                print(f"   📏 Text extracted: {content_length} characters")
                
                if content_length < 100:
                    print("   ❗ LIMITED TEXT - May be partially image-based")
                else:
                    print("   ✅ GOOD TEXT EXTRACTION")
            else:
                print("   ❌ IMAGE-BASED PDF - No text extracted")
            
            # Check for manual data
            if file_path.name in self.manual_data:
                data = self.manual_data[file_path.name]
                print(f"   📝 MANUAL DATA AVAILABLE: {len(data)} items")
                # Show key data points
                for key in ['gross_salary', 'net_salary', 'basic_salary']:
                    if key in data:
                        display_key = key.replace('_', ' ').title()
                        print(f"     - {display_key}: ₹{data[key]:,.2f}")
            else:
                print("   💡 SUGGESTION: Type 'manual' to enter data for this document")

    def show_help(self):
        """Show available commands in interactive mode."""
        print("\n📋 AVAILABLE COMMANDS:")
        print("  - Ask any question about your documents")
        print("  - 'manual' - Enter salary data manually for image-based PDFs")
        print("  - 'status' - Show detailed document status")
        print("  - 'export' - Save your queries and responses to a file")
        print("  - 'exit' - Exit interactive mode")
        print("  - 'help' - Show this help message")
        print("\n💡 FOR IMAGE-BASED PDFs:")
        print("  - Type 'manual' to enter salary data")
        print("  - The system will use your manual data for calculations")
        print("  - You can still ask questions about the manually entered data")

    def handle_export(self, query_history):
        """Handle export of query history."""
        if not query_history:
            print("❌ No queries to export")
            return
        
        print("\n📊 EXPORT OPTIONS:")
        print("1. PDF (Single query)")
        print("2. CSV (All queries)")
        print("3. JSON (All queries)")
        
        choice = input("Select format (1-3): ").strip()
        if choice == '1' and query_history:
            latest_query = query_history[-1]
            self.export_to_pdf(latest_query['query'], latest_query['response'])
        elif choice == '2':
            self.export_to_csv(query_history)
        elif choice == '3':
            self.export_to_json({
                'queries': query_history,
                'exported_at': datetime.now().isoformat(),
                'total_queries': len(query_history)
            })
        else:
            print("❌ Invalid option")

    def export_to_pdf(self, query: str, response: str, filename: str = None):
        """Export query and response to PDF format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Document Analysis Report", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, txt=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        if self.current_user:
            pdf.cell(200, 10, txt=f"User: {self.current_user['username']}", ln=True)
        pdf.ln(10)
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt="Query:", ln=True)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 8, txt=query)
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt="Response:", ln=True)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 8, txt=response)
        export_path = self.output_folder / filename
        pdf.output(str(export_path))
        print(f"📄 PDF export saved: {export_path}")

    def export_to_csv(self, queries: List[Dict], filename: str = None):
        """Export multiple queries and responses to CSV format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.csv"
        export_path = self.output_folder / filename
        try:
            with open(export_path, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['timestamp', 'query', 'response', 'user']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for query_data in queries:
                    writer.writerow({
                        'timestamp': query_data.get('timestamp', datetime.now().isoformat()),
                        'query': query_data.get('query', ''),
                        'response': query_data.get('response', ''),
                        'user': self.current_user['username'] if self.current_user else 'Anonymous'
                    })
            print(f"📊 CSV export saved: {export_path}")
        except Exception as e:
            print(f"❌ CSV export failed: {e}")

    def export_to_json(self, data: Dict, filename: str = None):
        """Export data to JSON format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.json"
        export_path = self.output_folder / filename
        try:
            with open(export_path, 'w', encoding='utf-8') as jsonfile:
                json.dump(data, jsonfile, indent=2, ensure_ascii=False)
            print(f"📋 JSON export saved: {export_path}")
        except Exception as e:
            print(f"❌ JSON export failed: {e}")

    def api_query_mode(self):
        """Run in API query mode."""
        print("🚀 Starting API Query Mode...")
        print("📡 This mode is designed for external applications")
        print("ℹ️  Use interactive mode for direct queries")
        self.interactive_mode()

    def cleanup(self):
        """Clean up resources and end session."""
        if self.current_user:
            self.logout_user()
        if hasattr(self, 'conn'):
            self.conn.close()
        print("✅ Cleanup completed!")

    def run(self):
        """Main execution flow with all advanced features."""
        try:
            if not self.authenticate_user():
                return
            print("\n📦 Processing documents...")
            file_ids = self.process_documents_in_batches(batch_size=8)
            if not file_ids:
                print("❌ No documents were processed successfully")
                return
            self.setup_assistant(file_ids)
            self.interactive_mode()
        except KeyboardInterrupt:
            print("\n\n⚠️  Operation cancelled by user")
        except Exception as e:
            print(f"\n❌ Fatal error: {e}")
            import traceback
            traceback.print_exc()
        finally:
            self.cleanup()

# =============================================================================
# MAIN EXECUTION
# =============================================================================

if __name__ == "__main__":
    import sys
    assistant = AdvancedPDFAnalysisAssistant()
    if len(sys.argv) > 1:
        if sys.argv[1] == "--api-mode":
            assistant.api_query_mode()
        else:
            print("Usage: python main.py")
    else:
        assistant.run()