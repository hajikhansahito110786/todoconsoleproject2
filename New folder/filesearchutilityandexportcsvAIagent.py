from openai import OpenAI
from dotenv import load_dotenv
import os
import time
import warnings
import glob
import json
import re
import csv
import PyPDF2
import docx
import sqlite3
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
import pandas as pd
from fpdf import FPDF
import uuid
import openpyxl
# =============================================================================
# CONFIGURATION AND INITIALIZATION
# =============================================================================

warnings.filterwarnings("ignore", category=DeprecationWarning)

class AdvancedPDFAnalysisAssistant:
    """
    Advanced document analysis system with multi-format support, batch processing,
    export capabilities, and user authentication.
    """
    
    def __init__(self):
        """Initialize the advanced document analysis system."""
        load_dotenv(dotenv_path="/root/filesearchfolder/.env")
        self.client = OpenAI()
        self.assistant_id = None
        self.vector_store_id = None
        self.file_ids = []
        self.current_user = None
        self.thread_id = None
        self.initialize_environment()
        self.init_database()

    def initialize_environment(self):
        """Initialize required directories and validate environment."""
        self.base_dir = Path.cwd()
        self.docs_folder = self.base_dir / "documents"
        self.output_folder = self.base_dir / "output"
        self.db_folder = self.base_dir / "database"
        self.temp_folder = self.base_dir / "temp"
        
        # Create necessary directories
        for folder in [self.docs_folder, self.output_folder, self.db_folder, self.temp_folder]:
            folder.mkdir(exist_ok=True)
        
        # Validate API key
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY not found in environment variables")

    def init_database(self):
        """Initialize SQLite database for user authentication and session tracking."""
        self.db_path = self.db_folder / "document_analysis.db"
        self.conn = sqlite3.connect(self.db_path)
        self.cursor = self.conn.cursor()
        
        # Create users table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                email TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                last_login TIMESTAMP
            )
        ''')
        
        # Create sessions table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                session_id TEXT UNIQUE,
                start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                end_time TIMESTAMP,
                queries_executed INTEGER DEFAULT 0,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')
        
        # Create query_history table
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS query_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT,
                query_text TEXT,
                response_text TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (session_id) REFERENCES sessions (session_id)
            )
        ''')
        
        self.conn.commit()

    def hash_password(self, password: str) -> str:
        """Hash password using SHA-256."""
        return hashlib.sha256(password.encode()).hexdigest()

    def register_user(self, username: str, password: str, email: str = None) -> bool:
        """Register a new user."""
        try:
            password_hash = self.hash_password(password)
            self.cursor.execute(
                "INSERT INTO users (username, password_hash, email) VALUES (?, ?, ?)",
                (username, password_hash, email)
            )
            self.conn.commit()
            print(f"✅ User '{username}' registered successfully!")
            return True
        except sqlite3.IntegrityError:
            print("❌ Username already exists")
            return False
        except Exception as e:
            print(f"❌ Registration failed: {e}")
            return False

    def login_user(self, username: str, password: str) -> bool:
        """Authenticate user and start session."""
        try:
            password_hash = self.hash_password(password)
            self.cursor.execute(
                "SELECT id FROM users WHERE username = ? AND password_hash = ?",
                (username, password_hash)
            )
            user = self.cursor.fetchone()
            
            if user:
                self.current_user = {
                    'id': user[0],
                    'username': username
                }
                
                session_id = str(uuid.uuid4())
                self.cursor.execute(
                    "INSERT INTO sessions (user_id, session_id) VALUES (?, ?)",
                    (user[0], session_id)
                )
                self.conn.commit()
                
                self.current_user['session_id'] = session_id
                print(f"✅ Welcome back, {username}!")
                return True
            else:
                print("❌ Invalid username or password")
                return False
                
        except Exception as e:
            print(f"❌ Login failed: {e}")
            return False

    def logout_user(self):
        """End current user session."""
        if self.current_user and 'session_id' in self.current_user:
            self.cursor.execute(
                "UPDATE sessions SET end_time = CURRENT_TIMESTAMP WHERE session_id = ?",
                (self.current_user['session_id'],)
            )
            self.conn.commit()
            print(f"👋 Goodbye, {self.current_user['username']}!")
        self.current_user = None

    def get_supported_files(self, folder_path: Path) -> List[Path]:
        """Get all supported document files."""
        all_files = list(folder_path.iterdir())
        supported_extensions = ['.pdf', '.docx', '.txt', '.csv', '.xlsx']
        files = [
            file for file in all_files 
            if file.is_file() and file.suffix.lower() in supported_extensions
        ]
        return self._validate_files(files)

    def _validate_files(self, files: List[Path]) -> List[Path]:
        """Validate files for size and format requirements."""
        valid_files = []
        max_size_mb = 20
        
        for file_path in files:
            try:
                file_size_mb = file_path.stat().st_size / (1024 * 1024)
                if file_size_mb > max_size_mb:
                    print(f"⚠️  Skipping {file_path.name} - Too large ({file_size_mb:.1f}MB)")
                    continue
                valid_files.append(file_path)
            except Exception as e:
                print(f"⚠️  Skipping {file_path.name} - Error: {e}")
        return valid_files

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file formats."""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_path.suffix.lower() in ['.txt', '.csv']:
                return self._extract_text_from_text_file(file_path)
            elif file_path.suffix.lower() == '.xlsx':
                return self._extract_text_from_excel(file_path)
            else:
                return ""
        except Exception as e:
            print(f"❌ Error extracting text from {file_path.name}: {e}")
            return ""

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"❌ PDF extraction error: {e}")
        return text

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            return ""

    def _extract_text_from_text_file(self, file_path: Path) -> str:
        """Extract text from text-based files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"❌ Text file extraction error: {e}")
            return ""

    def _extract_text_from_excel(self, file_path: Path) -> str:
        """Extract text from Excel files."""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            print(f"❌ Excel extraction error: {e}")
            return ""

    def process_documents_in_batches(self, batch_size: int = 10) -> List[str]:
        """Process documents in batches."""
        all_files = self.get_supported_files(self.docs_folder)
        if not all_files:
            print("ℹ️  No documents found for processing")
            return []
        
        print(f"📦 Found {len(all_files)} documents. Processing in batches of {batch_size}...")
        processed_files = []
        for i in range(0, len(all_files), batch_size):
            batch = all_files[i:i + batch_size]
            print(f"\n🔄 Processing batch {i//batch_size + 1}/{(len(all_files)-1)//batch_size + 1}")
            try:
                batch_file_ids = self._process_batch(batch)
                processed_files.extend(batch_file_ids)
                print(f"✅ Batch processed successfully. Total files: {len(processed_files)}")
                if i + batch_size < len(all_files):
                    time.sleep(5)
            except Exception as e:
                print(f"❌ Batch processing failed: {e}")
                continue
        return processed_files

    def _process_batch(self, batch_files: List[Path]) -> List[str]:
        """Process a single batch of files."""
        batch_file_ids = []
        processed_filenames = set()
        
        for file_path in batch_files:
            try:
                if file_path.name.lower() in processed_filenames:
                    print(f"⚠️  Skipping duplicate: {file_path.name}")
                    continue
                    
                if file_path.suffix.lower() != '.pdf':
                    text_content = self.extract_text_from_file(file_path)
                    if text_content:
                        temp_file = self.temp_folder / f"{file_path.stem}_converted.txt"
                        with open(temp_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                        with open(temp_file, 'rb') as f:
                            file_obj = self.client.files.create(file=f, purpose="assistants")
                            batch_file_ids.append(file_obj.id)
                            processed_filenames.add(file_path.name.lower())
                            print(f"✅ Uploaded converted: {file_path.name}")
                        temp_file.unlink()
                    else:
                        print(f"⚠️  Skipping {file_path.name} - No text content extracted")
                else:
                    with open(file_path, 'rb') as f:
                        file_obj = self.client.files.create(file=f, purpose="assistants")
                        batch_file_ids.append(file_obj.id)
                        processed_filenames.add(file_path.name.lower())
                        print(f"✅ Uploaded: {file_path.name}")
            except Exception as e:
                print(f"❌ Failed to process {file_path.name}: {e}")
        return batch_file_ids

    def wait_for_vector_store_processing(self, vector_store_id: str, max_attempts: int = 30, delay: int = 5):
        """Wait for vector store processing to complete."""
        print("⏳ Waiting for vector store processing...")
        for attempt in range(max_attempts):
            try:
                vector_store = self.client.vector_stores.retrieve(vector_store_id)
                status = vector_store.status
                if status == "completed":
                    print("✅ Vector store processing completed!")
                    return True
                elif status == "failed":
                    print("❌ Vector store processing failed!")
                    return False
                elif status in ["in_progress", "cancelling"]:
                    print(f"🔄 Processing... (Attempt {attempt + 1}/{max_attempts})")
                    time.sleep(delay)
                else:
                    print(f"ℹ️  Unknown vector store status: {status}")
                    time.sleep(delay)
            except Exception as e:
                print(f"⚠️  Error checking vector store status: {e}")
                time.sleep(delay)
        print("❌ Vector store processing timed out!")
        return False

    def setup_assistant(self, file_ids: List[str]):
        """Set up the assistant with the provided file IDs."""
        try:
            print("🏗️  Creating vector store...")
            vector_store = self.client.vector_stores.create(
                name=f"Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                file_ids=file_ids
            )
            self.vector_store_id = vector_store.id
            self.wait_for_vector_store_processing(self.vector_store_id)
            
            print("👨‍💼 Creating assistant...")
            assistant = self.client.beta.assistants.create(
                name="Advanced Document Analysis Assistant",
                instructions="""You are an expert document analysis assistant with multi-format support.
                Search across ALL uploaded documents to find relevant information.
                If the selected file doesn't contain the answer, search in other documents.
                Provide comprehensive answers based on all available documents.
                When answering, mention which documents contained the relevant information.
                If you find conflicting information across documents, mention the differences.
                Provide specific references to document names when possible.""",
                model="gpt-4o",
                tools=[{"type": "file_search"}],
                tool_resources={"file_search": {"vector_store_ids": [self.vector_store_id]}}
            )
            self.assistant_id = assistant.id
            
            thread = self.client.beta.threads.create()
            self.thread_id = thread.id
            print("✅ Assistant setup completed successfully!")
        except Exception as e:
            print(f"❌ Assistant setup failed: {e}")
            raise

    def execute_query(self, query: str):
        """Execute a query against the assistant."""
        print(f"🤔 Asking: {query}")
        try:
            self.client.beta.threads.messages.create(
                thread_id=self.thread_id,
                role="user",
                content=query
            )
            
            run = self.client.beta.threads.runs.create(
                thread_id=self.thread_id,
                assistant_id=self.assistant_id
            )
            
            print("⏳ Processing your query...")
            while run.status in ['queued', 'in_progress']:
                time.sleep(2)
                run = self.client.beta.threads.runs.retrieve(
                    thread_id=self.thread_id,
                    run_id=run.id
                )
            
            if run.status == 'completed':
                messages = self.client.beta.threads.messages.list(thread_id=self.thread_id)
                assistant_messages = [msg for msg in messages.data if msg.role == 'assistant']
                
                if assistant_messages:
                    response = assistant_messages[0].content[0].text.value
                    print(f"🤖 Assistant: {response}")
                    
                    if self.current_user and 'session_id' in self.current_user:
                        self.cursor.execute(
                            "INSERT INTO query_history (session_id, query_text, response_text) VALUES (?, ?, ?)",
                            (self.current_user['session_id'], query, response)
                        )
                        self.conn.commit()
                    return response
                else:
                    print("❌ No response from assistant")
                    return None
            else:
                print(f"❌ Run failed with status: {run.status}")
                if run.last_error:
                    print(f"Error: {run.last_error.message}")
                return None
        except Exception as e:
            print(f"❌ Error executing query: {e}")
            return None

    def analyze_excel_files(self, file1_path: str, file2_path: str, query: str) -> str:
        """Specialized method to analyze and compare Excel files."""
        try:
            df1 = pd.read_excel(file1_path)
            df2 = pd.read_excel(file2_path)
            
            print(f"📊 File 1 shape: {df1.shape}, File 2 shape: {df2.shape}")
            print(f"📊 File 1 columns: {list(df1.columns)}")
            print(f"📊 File 2 columns: {list(df2.columns)}")
            
            result = f"""
## Excel Files Analysis

### File 1: {Path(file1_path).name}
- **Rows:** {len(df1)}
- **Columns:** {len(df1.columns)}
- **Column Names:** {list(df1.columns)}

### File 2: {Path(file2_path).name}  
- **Rows:** {len(df2)}
- **Columns:** {len(df2.columns)}
- **Column Names:** {list(df2.columns)}
"""
            common_cols = set(df1.columns) & set(df2.columns)
            if common_cols:
                result += f"\n### Common Columns: {list(common_cols)}\n"
                for col in common_cols:
                    result += f"\n#### Matching analysis on column: '{col}'\n"
                    values1 = set(df1[col].dropna().astype(str))
                    values2 = set(df2[col].dropna().astype(str))
                    common_values = values1 & values2
                    result += f"- Common values: {len(common_values)}\n"
                    
                    if common_values:
                        sample_values = list(common_values)[:5]
                        result += f"- Sample common values: {sample_values}\n"
                        matches_df1 = df1[df1[col].astype(str).isin(common_values)]
                        matches_df2 = df2[df2[col].astype(str).isin(common_values)]
                        result += f"- Rows in File 1 with matching values: {len(matches_df1)}\n"
                        result += f"- Rows in File 2 with matching values: {len(matches_df2)}\n"
                        
                        if len(matches_df1) > 0 and len(matches_df2) > 0:
                            result += "\n**Sample matching rows:**\n"
                            result += f"\nFile 1 - First 3 matching rows:\n{matches_df1.head(3).to_string()}\n"
                            result += f"\nFile 2 - First 3 matching rows:\n{matches_df2.head(3).to_string()}\n"
            return result
        except Exception as e:
            return f"Error analyzing Excel files: {str(e)}"

    def _process_special_query(self, query: str, filename: str) -> str:
        """Process queries - always search across ALL documents"""
        try:
            query_lower = query.lower()
            
            # Check if this is an Excel file comparison query
            if any(x in query_lower for x in ['link', 'compare', 'matching', 'difference']) and any(x in query_lower for x in ['.xlsx', '.csv', 'excel']):
                import re
                files = re.findall(r'(\w+\.xlsx|\w+\.csv)', query)
                if len(files) >= 2:
                    file1_name, file2_name = files[0], files[1]
                    file1_path = self.docs_folder / file1_name
                    file2_path = self.docs_folder / file2_name
                    
                    if file1_path.exists() and file2_path.exists():
                        print(f"🔍 Performing Excel analysis: {file1_name} vs {file2_name}")
                        return self.analyze_excel_files(str(file1_path), str(file2_path), query)
                    else:
                        return f"One or both files not found: {file1_name}, {file2_name}"
            
            # ALWAYS search across ALL documents regardless of selected file
            print(f"🔍 Searching across ALL documents for: {query}")
            
            # Get list of all documents for context
            all_files = self.get_supported_files(self.docs_folder)
            file_list = [f.name for f in all_files]
            
            enhanced_query = f"""
Please search across ALL available documents to answer this query: "{query}"

Available documents: {', '.join(file_list)}

Search through all these documents comprehensively to find the most relevant information.
If the information exists in any of the documents, provide a complete answer.
If you find conflicting information across documents, mention the differences.
If no information is found in any document, please state that clearly.

Provide specific references to which document contained the information when possible.
"""
            response = self.execute_query(enhanced_query)
            return response if response else "No response received from the assistant"
            
        except Exception as e:
            return f"Error processing query: {str(e)}"

    def api_query_mode(self):
        """Run in API query mode - maintains assistant state and processes files properly."""
        print("🚀 Starting API Query Mode - Ready for requests...")
        print("📡 Monitoring signal file: /tmp/assistant_request.signal")
        
        try:
            print("📦 Initial processing of all documents...")
            file_ids = self.process_documents_in_batches(batch_size=8)
            if file_ids:
                print("🏗️ Setting up assistant with all documents...")
                self.setup_assistant(file_ids)
                print("✅ Assistant initialized and ready for queries!")
                print(f"📊 Loaded {len(file_ids)} documents")
                
                # List all processed files
                all_files = self.get_supported_files(self.docs_folder)
                print("📋 Available documents for search:")
                for file in all_files:
                    print(f"   - {file.name}")
            else:
                print("⚠️ No documents found for processing")
        except Exception as e:
            print(f"❌ Initial setup failed: {e}")
        
        try:
            while True:
                if os.path.exists("/tmp/assistant_request.signal"):
                    try:
                        with open("/tmp/assistant_request.signal", 'r') as f:
                            request_file_path = f.read().strip()
                        print(f"📨 Received request: {request_file_path}")
                        os.unlink("/tmp/assistant_request.signal")
                        
                        if os.path.exists(request_file_path):
                            with open(request_file_path, 'r') as f:
                                request_data = json.load(f)
                            
                            action = request_data.get("action")
                            query = request_data.get("query")
                            filename = request_data.get("filename")
                            
                            if action == "query" and query and filename:
                                print(f"🎯 Processing query: '{query}'")
                                print(f"📁 Selected file (for reference): {filename}")
                                response = self._process_special_query(query, filename)
                                response_file = request_file_path + ".response"
                                with open(response_file, 'w') as f:
                                    json.dump({"response": response}, f)
                                print("✅ Response sent successfully")
                                print("=" * 50)
                    except Exception as e:
                        print(f"❌ Error processing request: {e}")
                        try:
                            if 'request_file_path' in locals():
                                response_file = request_file_path + ".response"
                                with open(response_file, 'w') as f:
                                    json.dump({"response": f"Error: {str(e)}"}, f)
                        except:
                            pass
                time.sleep(2)
        except KeyboardInterrupt:
            print("\n🛑 API Query Mode stopped by user")
        except Exception as e:
            print(f"❌ API Query Mode error: {e}")

    def interactive_mode(self):
        """Enter interactive mode for querying documents."""
        print("\n" + "="*60)
        print("💬 INTERACTIVE MODE - Ask questions about your documents")
        print("Type 'exit' to quit, 'export' to save results, or 'help' for options")
        print("="*60)
        
        query_count = 0
        query_history = []
        
        while True:
            user_input = input("\n🎯 Your question: ").strip()
            if user_input.lower() in ['exit', 'quit', 'bye']:
                print("Exiting interactive mode...")
                break
            elif user_input.lower() in ['help', '?']:
                self.show_help()
                continue
            elif user_input.lower() == 'export':
                self.handle_export(query_history)
                continue
            elif not user_input:
                continue
            
            response = self.execute_query(user_input)
            if response:
                query_count += 1
                query_history.append({
                    'query': user_input,
                    'response': response,
                    'timestamp': datetime.now().isoformat()
                })
                if query_count % 3 == 0:
                    print("\n💡 Tip: Type 'export' to save your results to a file")

    def show_help(self):
        """Show available commands in interactive mode."""
        print("\n📋 AVAILABLE COMMANDS:")
        print("  - Ask any question about your documents")
        print("  - 'export' - Save your queries and responses to a file")
        print("  - 'exit' - Exit interactive mode")
        print("  - 'help' - Show this help message")
        print("\n💡 EXAMPLE QUESTIONS:")
        print("  - 'What are the main topics in these documents?'")
        print("  - 'Summarize the key findings'")
        print("  - 'Find information about [specific topic]'")

    def handle_export(self, query_history):
        """Handle export of query history."""
        if not query_history:
            print("❌ No queries to export")
            return
        
        print("\n📊 EXPORT OPTIONS:")
        print("1. PDF (Single query)")
        print("2. CSV (All queries)")
        print("3. JSON (All queries)")
        print("4. XLSX (All queries)")
        print("5. DOCX (All queries)")
        
        choice = input("Select format (1-5): ").strip()
        if choice == '1' and query_history:
            latest_query = query_history[-1]
            self.export_to_pdf(latest_query['query'], latest_query['response'])
        elif choice == '2':
            self.export_to_csv(query_history)
        elif choice == '3':
            self.export_to_json({
                'queries': query_history,
                'exported_at': datetime.now().isoformat(),
                'total_queries': len(query_history)
            })
        elif choice == '4':
            self.export_to_xlsx(query_history)
        elif choice == '5':
            self.export_to_docx(query_history)
        else:
            print("❌ Invalid option")

    def export_to_pdf(self, query: str, response: str, filename: str = None):
        """Export query and response to PDF format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Document Analysis Report", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, txt=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        if self.current_user:
            pdf.cell(200, 10, txt=f"User: {self.current_user['username']}", ln=True)
        pdf.ln(10)
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt="Query:", ln=True)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 8, txt=query)
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt="Response:", ln=True)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 8, txt=response)
        export_path = self.output_folder / filename
        pdf.output(str(export_path))
        print(f"📄 PDF export saved: {export_path}")

    def export_to_csv(self, queries: List[Dict], filename: str = None):
        """Export multiple queries and responses to CSV format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.csv"
        export_path = self.output_folder / filename
        try:
            with open(export_path, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['timestamp', 'query', 'response', 'user']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                for query_data in queries:
                    writer.writerow({
                        'timestamp': query_data.get('timestamp', datetime.now().isoformat()),
                        'query': query_data.get('query', ''),
                        'response': query_data.get('response', ''),
                        'user': self.current_user['username'] if self.current_user else 'Anonymous'
                    })
            print(f"📊 CSV export saved: {export_path}")
        except Exception as e:
            print(f"❌ CSV export failed: {e}")

    def export_to_json(self, data: Dict, filename: str = None):
        """Export data to JSON format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.json"
        export_path = self.output_folder / filename
        try:
            with open(export_path, 'w', encoding='utf-8') as jsonfile:
                json.dump(data, jsonfile, indent=2, ensure_ascii=False)
            print(f"📋 JSON export saved: {export_path}")
        except Exception as e:
            print(f"❌ JSON export failed: {e}")

    def export_to_xlsx(self, queries: List[Dict], filename: str = None):
        """Export multiple queries and responses to XLSX format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.xlsx"
        export_path = self.output_folder / filename
        try:
            data = []
            for query_data in queries:
                data.append({
                    'Timestamp': query_data.get('timestamp', datetime.now().isoformat()),
                    'Query': query_data.get('query', ''),
                    'Response': query_data.get('response', ''),
                    'User': self.current_user['username'] if self.current_user else 'Anonymous'
                })
            df = pd.DataFrame(data)
            with pd.ExcelWriter(export_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Query History', index=False)
                worksheet = writer.sheets['Query History']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            print(f"📊 XLSX export saved: {export_path}")
        except Exception as e:
            print(f"❌ XLSX export failed: {e}")

    def export_to_docx(self, queries: List[Dict], filename: str = None):
        """Export multiple queries and responses to DOCX format."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.docx"
        export_path = self.output_folder / filename
        try:
            doc = docx.Document()
            doc.add_heading('Document Analysis Report', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if self.current_user:
                doc.add_paragraph(f"User: {self.current_user['username']}")
            for i, query_data in enumerate(queries, 1):
                doc.add_heading(f"Query {i}", level=2)
                doc.add_paragraph("Query:", style='Heading 3')
                doc.add_paragraph(query_data.get('query', ''))
                doc.add_paragraph("Response:", style='Heading 3')
                doc.add_paragraph(query_data.get('response', ''))
                if i < len(queries):
                    doc.add_paragraph("-" * 50)
            doc.save(str(export_path))
            print(f"📝 DOCX export saved: {export_path}")
        except Exception as e:
            print(f"❌ DOCX export failed: {e}")

    def run(self):
        """Main execution flow with all advanced features."""
        try:
            if not self.authenticate_user():
                return
            print("\n📦 Processing documents...")
            file_ids = self.process_documents_in_batches(batch_size=8)
            if not file_ids:
                print("❌ No documents were processed successfully")
                return
            self.setup_assistant(file_ids)
            predefined_query = """Just answer query of interactive mode"""
            self.execute_query(predefined_query)
            self.interactive_mode()
        except KeyboardInterrupt:
            print("\n\n⚠️  Operation cancelled by user")
        except Exception as e:
            print(f"\n❌ Fatal error: {e}")
            import traceback
            traceback.print_exc()
        finally:
            self.cleanup()

    def authenticate_user(self) -> bool:
        """Handle user authentication with multiple options."""
        print("🔐 Document Analysis System - Authentication")
        print("="*50)
        while True:
            print("\nOptions:")
            print("1. Login")
            print("2. Register")
            print("3. Continue as Guest")
            print("4. Exit")
            choice = input("Select option (1-4): ").strip()
            if choice == '1':
                username = input("Username: ").strip()
                password = input("Password: ").strip()
                if self.login_user(username, password):
                    return True
            elif choice == '2':
                username = input("New username: ").strip()
                password = input("New password: ").strip()
                email = input("Email (optional): ").strip() or None
                if self.register_user(username, password, email):
                    if self.login_user(username, password):
                        return True
            elif choice == '3':
                print("ℹ️  Continuing as guest user (limited functionality)")
                return True
            elif choice == '4':
                print("👋 Goodbye!")
                return False
            else:
                print("❌ Invalid option")

    def cleanup(self):
        """Clean up resources and end session."""
        if self.current_user:
            self.logout_user()
        if hasattr(self, 'conn'):
            self.conn.close()

# =============================================================================
# MAIN EXECUTION
# =============================================================================

if __name__ == "__main__":
    import sys
    assistant = AdvancedPDFAnalysisAssistant()
    if len(sys.argv) > 1:
        if sys.argv[1] == "--api-mode":
            assistant.api_query_mode()
        elif sys.argv[1] == "--simple-api":
            print("Simple API mode not available, using api-query-mode instead")
            assistant.api_query_mode()
        else:
            print("Usage:")
            print("  python filesearchutilityandexportcsvAIagent.py          # Interactive mode")
            print("  python filesearchutilityandexportcsvAIagent.py --api-mode  # API mode with pre-loaded docs")
    else:
        assistant.run()